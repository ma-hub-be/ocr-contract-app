from flask import Flask, render_template, request, redirect, url_for, send_file
from contract_auto import extract_text, preprocess_image
from pathlib import Path
from functools import wraps
import time
import PyPDF2
import os
import re

# ===== Basic認証 =====
def check_auth(username, password):
    return username == 'maika' and password == 'perogostini'

def authenticate():
    from flask import Response
    return Response(
        'ログインが必要です', 401,
        {'WWW-Authenticate': 'Basic realm="Login Required"'}
    )

def requires_auth(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        auth = request.authorization
        if not auth or not check_auth(auth.username, auth.password):
            return authenticate()
        return f(*args, **kwargs)
    return decorated

# ===== テキスト正規化（差分比較の精度向上・最終版） =====
def normalize_text(text):
    """OCR結果を正規化して比較しやすくする（最終版）"""
    # --- ページ X --- のヘッダーを除去（OCR由来）
    text = re.sub(r'-+\s*ページ\s*\d+\s*-+', '', text)

    # 全角スペース→除去
    text = text.replace('\u3000', '')

    # タブを除去
    text = text.replace('\t', '')

    # 半角スペースを除去
    text = text.replace(' ', '')

    # 改行を除去（全テキストを1つに結合）
    text = text.replace('\n', '')
    text = text.replace('\r', '')

    # 全角数字→半角数字
    zen_to_han = str.maketrans('０１２３４５６７８９', '0123456789')
    text = text.translate(zen_to_han)

    # 全角英字→半角英字
    zen_alpha = 'ＡＢＣＤＥＦＧＨＩＪＫＬＭＮＯＰＱＲＳＴＵＶＷＸＹＺａｂｃｄｅｆｇｈｉｊｋｌｍｎｏｐｑｒｓｔｕｖｗｘｙｚ'
    han_alpha = 'ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz'
    text = text.translate(str.maketrans(zen_alpha, han_alpha))

    # 全角記号の正規化
    text = text.replace('．', '.').replace('，', ',').replace('：', ':')
    text = text.replace('；', ';').replace('（', '(').replace('）', ')')

    # ===== 句点（。）で分割して1文=1行にする =====
    parts = re.split(r'(。)', text)

    sentences = []
    buffer = ''
    for part in parts:
        buffer += part
        if part == '。':
            sentences.append(buffer.strip())
            buffer = ''
    if buffer.strip():
        sentences.append(buffer.strip())

    # ===== 見出し行を分離 =====
    final_lines = []
    for sentence in sentences:
        # 「第X条(...)」のパターンを分離
        parts = re.split(r'(第\d+条[（(][^）)]*[）)])', sentence)
        for p in parts:
            p = p.strip()
            if p:
                final_lines.append(p)

    return '\n'.join(final_lines)

# ===== Flaskアプリ =====
app = Flask(__name__)
app.secret_key = 'your-secret-key-here'

# アップロードフォルダの設定
UPLOAD_FOLDER = Path('uploads')
UPLOAD_FOLDER.mkdir(exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB上限

def extract_text_from_file(filepath):
    """
    ファイルからテキストを抽出する共通関数
    PDF、画像、Word、Excelに対応
    """
    file_ext = Path(filepath).suffix.lower()

    if file_ext == '.pdf':
        return extract_text(str(filepath))

    elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp']:
        from PIL import Image
        import pytesseract

        image = Image.open(filepath)
        processed_image = preprocess_image(image)

        custom_config = r'--oem 3 --psm 6 -l jpn+eng -c preserve_interword_spaces=1'
        return pytesseract.image_to_string(processed_image, config=custom_config)

    elif file_ext == '.docx':
        from docx import Document

        doc = Document(filepath)
        extracted_text = ""

        for para in doc.paragraphs:
            extracted_text += para.text + "\n"

        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join([cell.text for cell in row.cells])
                extracted_text += row_text + "\n"

        return extracted_text

    elif file_ext in ['.xlsx', '.xls']:
        from openpyxl import load_workbook

        wb = load_workbook(filepath, data_only=True)
        extracted_text = ""

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            extracted_text += f"\n=== シート: {sheet_name} ===\n\n"

            for row in ws.iter_rows(values_only=True):
                row_text = "\t".join([str(cell) if cell is not None else "" for cell in row])
                if row_text.strip():
                    extracted_text += row_text + "\n"

        return extracted_text

    else:
        raise ValueError(f"未対応のファイル形式です: {file_ext}")

@app.route('/')
@requires_auth
def index():
    """アップロード画面"""
    return render_template('upload.html')

@app.route('/upload', methods=['POST'])
@requires_auth
def upload_file():
    """ファイルをアップロードしてOCR/テキスト抽出処理"""
    if 'file' not in request.files:
        return "ファイルが選択されていません", 400

    file = request.files['file']

    if file.filename == '':
        return "ファイルが選択されていません", 400

    allowed_extensions = ['.pdf', '.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp', '.docx', '.xlsx', '.xls']
    file_ext = Path(file.filename).suffix.lower()

    if file_ext not in allowed_extensions:
        return f"対応していないファイル形式です。対応形式: {', '.join(allowed_extensions)}", 400

    filename = f"{int(time.time())}_{file.filename}"
    filepath = app.config['UPLOAD_FOLDER'] / filename
    file.save(filepath)

    try:
        extracted_text = extract_text_from_file(filepath)

        if file_ext == '.pdf':
            try:
                with open(filepath, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    text_check = pdf_reader.pages[0].extract_text()
                    file_type = "テキストPDF" if len(text_check.strip()) > 100 else "画像PDF"
            except:
                file_type = "画像PDF"
        elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp']:
            file_type = f"画像ファイル ({file_ext.upper()})"
        elif file_ext == '.docx':
            file_type = "Word文書 (.docx)"
        elif file_ext in ['.xlsx', '.xls']:
            file_type = f"Excelファイル ({file_ext.upper()})"

        print(f"ファイル種類: {file_type}")
        print(f"抽出文字数: {len(extracted_text)}")
        print(f"最初の100文字: {extracted_text[:100] if len(extracted_text) > 0 else '(空)'}")

        Path("results").mkdir(exist_ok=True)
        result_file = Path("results") / f"抽出結果_{int(time.time())}.txt"

        with open(result_file, 'w', encoding='utf-8') as f:
            f.write(f"ファイル名: {file.filename}\n")
            f.write(f"ファイル種類: {file_type}\n")
            f.write(f"処理日時: {time.strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write("=" * 60 + "\n\n")
            f.write(extracted_text)

        print(f"結果ファイル: {result_file.name}")

        if filepath.exists():
            os.remove(filepath)

        return render_template('result.html',
                             text=extracted_text,
                             pdf_type=file_type,
                             filename=result_file.name)

    except Exception as e:
        if filepath.exists():
            os.remove(filepath)

        print(f"エラー詳細: {str(e)}")
        import traceback
        traceback.print_exc()

        return f"処理エラー: {str(e)}", 500

@app.route('/download/<filename>')
@requires_auth
def download_file(filename):
    """結果ファイルをダウンロード"""
    file_path = Path("results") / filename
    if file_path.exists():
        return send_file(file_path, as_attachment=True)
    return "ファイルが見つかりません", 404

@app.route('/compare')
@requires_auth
def compare_page():
    """比較画面"""
    return render_template('compare.html')

@app.route('/compare_upload', methods=['POST'])
@requires_auth
def compare_upload():
    """2つのファイルを比較（全ファイル形式対応）"""
    if 'file1' not in request.files or 'file2' not in request.files:
        return "2つのファイルを選択してください", 400

    file1 = request.files['file1']
    file2 = request.files['file2']

    if file1.filename == '' or file2.filename == '':
        return "2つのファイルを選択してください", 400

    filepath1 = app.config['UPLOAD_FOLDER'] / f"{int(time.time())}_1_{file1.filename}"
    filepath2 = app.config['UPLOAD_FOLDER'] / f"{int(time.time())}_2_{file2.filename}"
    file1.save(filepath1)
    file2.save(filepath2)

    try:
        text1 = extract_text_from_file(filepath1)
        text2 = extract_text_from_file(filepath2)

        # テキスト正規化（段落ベースで統一）
        text1 = normalize_text(text1)
        text2 = normalize_text(text2)

        lines1 = text1.splitlines()
        lines2 = text2.splitlines()

        import difflib
        matcher = difflib.SequenceMatcher(None, lines1, lines2)

        highlighted_lines1 = []
        highlighted_lines2 = []

        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'equal':
                for i in range(i1, i2):
                    highlighted_lines1.append({
                        'type': 'normal',
                        'html': lines1[i]
                    })
                for j in range(j1, j2):
                    highlighted_lines2.append({
                        'type': 'normal',
                        'html': lines2[j]
                    })

            elif tag == 'delete':
                for i in range(i1, i2):
                    highlighted_lines1.append({
                        'type': 'delete',
                        'html': lines1[i]
                    })

            elif tag == 'insert':
                for j in range(j1, j2):
                    highlighted_lines2.append({
                        'type': 'insert',
                        'html': lines2[j]
                    })

            elif tag == 'replace':
                old_lines = lines1[i1:i2]
                new_lines = lines2[j1:j2]

                for idx in range(max(len(old_lines), len(new_lines))):
                    old_line = old_lines[idx] if idx < len(old_lines) else ""
                    new_line = new_lines[idx] if idx < len(new_lines) else ""

                    char_matcher = difflib.SequenceMatcher(None, old_line, new_line)

                    old_html = ""
                    new_html = ""

                    for ctag, ci1, ci2, cj1, cj2 in char_matcher.get_opcodes():
                        if ctag == 'equal':
                            old_html += old_line[ci1:ci2]
                            new_html += new_line[cj1:cj2]
                        elif ctag == 'delete':
                            old_html += f'<span class="char-delete">{old_line[ci1:ci2]}</span>'
                        elif ctag == 'insert':
                            new_html += f'<span class="char-insert">{new_line[cj1:cj2]}</span>'
                        elif ctag == 'replace':
                            old_html += f'<span class="char-change">{old_line[ci1:ci2]}</span>'
                            new_html += f'<span class="char-change">{new_line[cj1:cj2]}</span>'

                    if old_html:
                        highlighted_lines1.append({
                            'type': 'change',
                            'html': old_html
                        })

                    if new_html:
                        highlighted_lines2.append({
                            'type': 'change',
                            'html': new_html
                        })

        return render_template('compare_result_char.html',
                             file1_name=file1.filename,
                             file2_name=file2.filename,
                             highlighted_lines1=highlighted_lines1,
                             highlighted_lines2=highlighted_lines2)

    except Exception as e:
        print(f"比較エラー: {str(e)}")
        import traceback
        traceback.print_exc()
        return f"比較処理エラー: {str(e)}", 500

    finally:
        if filepath1.exists():
            os.remove(filepath1)
        if filepath2.exists():
            os.remove(filepath2)

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 10000))
    app.run(host="0.0.0.0", port=port)
